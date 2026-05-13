"""Generate JaneERP SOP as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ---------------------------------------------------------------------------
# Page margins
# ---------------------------------------------------------------------------
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
styles = doc.styles

def set_style(style_name, font_name, font_size, bold=False, color=None, space_before=None, space_after=None):
    try:
        s = styles[style_name]
    except KeyError:
        return
    s.font.name = font_name
    s.font.size = Pt(font_size)
    s.font.bold = bold
    if color:
        s.font.color.rgb = RGBColor(*color)
    pf = s.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)

set_style('Normal',    'Calibri', 11, space_after=4)
set_style('Heading 1', 'Calibri', 18, bold=True, color=(31, 73, 125),  space_before=18, space_after=6)
set_style('Heading 2', 'Calibri', 14, bold=True, color=(68, 114, 196), space_before=14, space_after=4)
set_style('Heading 3', 'Calibri', 12, bold=True, color=(68, 114, 196), space_before=10, space_after=3)

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(text='', bold=False, italic=False, style='Normal'):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    _add_inline(p, text)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    _add_inline(p, text)
    return p

def _add_inline(p, text):
    """Add text with **bold** and `code` inline markers."""
    import re
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            p.add_run(part)

def add_note(text):
    """Shaded note/tip box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # light blue shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'DCE6F1')
    pPr.append(shd)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 73, 125)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # header fill
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        fill = 'EEF3FB' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = ''
            p = row_cells[c_idx].paragraphs[0]
            _add_inline(p, str(cell_text))
            p.runs[0].font.size = Pt(10) if p.runs else None
            # alternating row colour
            tc = row_cells[c_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    # Column widths
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])

    doc.add_paragraph()  # spacing after table
    return table

def page_break():
    doc.add_page_break()

# ===========================================================================
# TITLE PAGE
# ===========================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('JaneERP')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(31, 73, 125)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Standard Operating Procedure')
run2.font.size = Pt(22)
run2.font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version 1.0  |  May 2026\nApplies To: All JaneERP Users').font.size = Pt(12)

page_break()

# ===========================================================================
# TABLE OF CONTENTS (manual)
# ===========================================================================
add_heading('Table of Contents', 1)
toc_items = [
    ('1.', 'Overview'),
    ('2.', 'Getting Started'),
    ('3.', 'User Roles & Permissions'),
    ('4.', 'Main Menu & Navigation'),
    ('5.', 'Sales & Orders'),
    ('6.', 'Purchase Orders & Receiving'),
    ('7.', 'Inventory Management'),
    ('8.', 'Products & Parts'),
    ('9.', 'Manufacturing'),
    ('10.', 'Reporting & Analytics'),
    ('11.', 'Task Management'),
    ('12.', 'System Administration'),
    ('13.', 'Data Import & Export'),
    ('14.', 'Troubleshooting & FAQs'),
]
for num, title_text in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(31, 73, 125)
    p.add_run(title_text)

page_break()

# ===========================================================================
# SECTION 1 — OVERVIEW
# ===========================================================================
add_heading('1. Overview', 1)
add_para('JaneERP is the company\'s internal management system for day-to-day operations. It is a Windows desktop application that brings together all key business functions into a single, integrated platform.')
add_para()
add_para('The system handles:')
add_bullet('**Sales orders** — manual orders and Shopify store orders')
add_bullet('**Purchase orders** — buying stock from suppliers')
add_bullet('**Inventory** — tracking stock quantities across multiple locations')
add_bullet('**Manufacturing** — tracking production work orders and costs')
add_bullet('**Reporting** — financial and operational reports')
add_bullet('**Team tasks** — assigning and tracking internal tasks')
add_para()
add_para('All users log in with a personal account. What you can see and do depends on your assigned role.')

page_break()

# ===========================================================================
# SECTION 2 — GETTING STARTED
# ===========================================================================
add_heading('2. Getting Started', 1)

add_heading('2.1  Logging In', 2)
add_numbered('Open the JaneERP application.')
add_numbered('Enter your **Username** and **Password**.')
add_numbered('Click **Login**.')
add_para()
add_note('Tip: If "Remember Username" is enabled, your username will be pre-filled next time.')
add_para()
add_para('**If your account is locked:** Contact your administrator. Accounts lock after a set number of failed login attempts and unlock automatically after a cooldown period (typically 15 minutes).')
add_para()
add_para('**First time using the system:** Your administrator will create your account and provide your credentials. Change your password on first login if prompted.')

add_heading('2.2  Session Timeout', 2)
add_para('The system automatically logs you out after **30 minutes of inactivity** to protect your account. You will be returned to the login screen. Any unsaved work may be lost — save frequently.')

add_heading('2.3  Logging Out', 2)
add_para('Close the main menu or use the logout option to end your session. Your login and logout times are recorded in the system audit log.')

page_break()

# ===========================================================================
# SECTION 3 — ROLES & PERMISSIONS
# ===========================================================================
add_heading('3. User Roles & Permissions', 1)
add_para('There are three roles in JaneERP. Your administrator assigns your role when they create your account.')
add_para()
add_table(
    ['Role', 'What You Can Do'],
    [
        ['Admin', 'Full access: user management, settings, all modules, all data'],
        ['Editor', 'Create and edit orders, inventory, manufacturing, tasks, reports'],
        ['Viewer', 'Read-only: dashboards, reports, product search'],
    ],
    col_widths=[1.5, 4.5]
)

add_para('Granular permissions can also be set per user for specific areas:')
add_para()
add_table(
    ['Permission', 'Controls Access To'],
    [
        ['Inventory', 'Inventory adjustment, stock transfers, location management'],
        ['Sales Orders', 'Creating and editing sales orders'],
        ['Manufacturing', 'Work orders and manufacturing orders'],
        ['Parts', 'Product and parts management'],
        ['Cycle Count', 'Physical inventory counting'],
        ['Tasks', 'Creating and managing team tasks'],
        ['Logs', 'Viewing audit and login logs'],
    ],
    col_widths=[1.8, 4.2]
)
add_note('If a button or menu item is greyed out or not visible, you likely do not have permission for that function. Contact your administrator.')

page_break()

# ===========================================================================
# SECTION 4 — MAIN MENU
# ===========================================================================
add_heading('4. Main Menu & Navigation', 1)
add_para('After logging in, you will see the Main Menu — the hub for all areas of the system.')

add_heading('4.1  Main Menu Sections', 2)
add_table(
    ['Section', 'Buttons / Functions'],
    [
        ['Sales', 'Sales Dashboard, Create Order, Customers'],
        ['Purchasing', 'Purchase Orders, Vendors, Reorder Report'],
        ['Inventory', 'Inventory Snapshot, Cycle Count, Adjust Stock, Stock Transfer, Locations'],
        ['Products', 'Products, Parts, Product Types, Attributes'],
        ['Manufacturing', 'Manufacturing Dashboard, Work Orders'],
        ['Analytics', 'KPI Dashboard, Reports, Breakeven Calculator'],
        ['Team', 'Task Manager, Activity Log'],
        ['Admin', 'User Management, Settings, Login Log'],
    ],
    col_widths=[1.8, 4.2]
)

add_heading('4.2  Notification Badges', 2)
add_para('You may see small numbered badges on certain buttons:')
add_bullet('**Tasks badge** — number of unread mentions in task comments')
add_bullet('**Unverified Items badge** — products added but not yet verified')
add_bullet('**Cycle Count badge** — locations overdue for a physical count')
add_para()
add_para('These badges are alerts that need your attention.')

page_break()

# ===========================================================================
# SECTION 5 — SALES & ORDERS
# ===========================================================================
add_heading('5. Sales & Orders', 1)

add_heading('5.1  Viewing Orders (Sales Dashboard)', 2)
add_numbered('From the Main Menu, click **Sales Dashboard**.')
add_numbered('Use the **Store** dropdown to filter by Shopify store or view all orders.')
add_numbered('Set a **date range** using the From/To date pickers.')
add_numbered('Use the **minimum amount** filter to narrow results by order value.')
add_numbered('Click **Fetch Latest** to pull new orders from Shopify.')
add_numbered('Click **Sync to ERP** to import Shopify orders into the database.')
add_para()
add_para('The order list shows order number, customer, date, total, and status.')

add_heading('5.2  Creating a Manual Sales Order', 2)
add_para('Use this for orders placed directly (not through Shopify).')
add_para()
add_numbered('From the Main Menu, click **Create Order**.')
add_numbered('**Search for the customer** — type part of their name or email and press Enter or click Search. If the customer does not exist, enter their details manually.')
add_numbered('Set the **Order Date**.')
add_numbered('Select the **Currency** (e.g., ZAR, USD). All prices entered will be in this currency.')
add_numbered('Set the **Order Type** (e.g., Wholesale, Retail).')
add_numbered('**Add line items:** click Add Item, search for a product by SKU or name, enter the quantity and unit price. Repeat for each product.')
add_numbered('**Apply a discount** (optional) — the customer\'s tier may apply automatically, or you can enter a fixed or percentage discount.')
add_numbered('Enter **shipping cost** if applicable.')
add_numbered('Review the **order total** (shown in both the selected currency and home currency).')
add_numbered('Set the **status** (Draft or Live).')
add_numbered('Click **Save**.')
add_para()
add_note('Draft orders are saved but not finalised. Set to Live once confirmed.')

add_heading('5.3  Customer Discount Tiers', 2)
add_para('Discount tiers allow automatic discounts for specific customers (e.g., wholesale accounts).')
add_bullet('Tiers are set up by Admins in Settings > Customer Tiers.')
add_bullet('When a customer with a tier is selected, the discount applies automatically.')
add_bullet('You can override the discount on any individual order.')

add_heading('5.4  Multi-Currency Orders', 2)
add_bullet('Select the correct currency when creating the order.')
add_bullet('Exchange rates are configured by your administrator in Settings.')
add_bullet('The system displays the equivalent home currency total for reference.')
add_bullet('Reports convert all orders back to home currency.')

page_break()

# ===========================================================================
# SECTION 6 — PURCHASE ORDERS
# ===========================================================================
add_heading('6. Purchase Orders & Receiving', 1)

add_heading('6.1  Creating a Purchase Order (PO)', 2)
add_numbered('From the Main Menu, click **Purchase Orders**.')
add_numbered('Click **New PO**.')
add_numbered('Select the **Supplier** from the dropdown.')
add_numbered('Set the **Expected Delivery Date**.')
add_numbered('Add line items: search for the product or part, enter quantity and unit cost.')
add_numbered('Review the totals.')
add_numbered('Click **Save**. The PO is saved as **Draft**.')

add_heading('6.2  PO Statuses', 2)
add_table(
    ['Status', 'Meaning'],
    [
        ['Draft', 'PO created but not yet sent to the supplier'],
        ['Sent', 'PO has been sent to the supplier'],
        ['Partially Received', 'Some items received, others still outstanding'],
        ['Received', 'All items received and inventory updated'],
        ['Cancelled', 'PO was cancelled'],
    ],
    col_widths=[2.0, 4.0]
)
add_para('Update the PO status manually as it progresses (e.g., mark as Sent once you have contacted the supplier).')

add_heading('6.3  Receiving Items', 2)
add_para('When goods arrive from a supplier:')
add_para()
add_numbered('Open the relevant PO in **Purchase Orders**.')
add_numbered('Click **Receive Items**.')
add_numbered('Enter the **actual quantities received** for each line.')
add_numbered('Select the **destination location** (where the stock will go).')
add_numbered('Click **Confirm Receipt**.')
add_para()
add_para('The system will update inventory quantities, record the transaction in the inventory log, and update the PO status to Partially Received or Received.')
add_para()
add_note('If you receive fewer items than ordered, the PO becomes Partially Received. Receive the remaining items when they arrive.')

add_heading('6.4  Auto-Reorder Report', 2)
add_numbered('Click **Reorder Report** from the Main Menu.')
add_numbered('Review the list of items below their reorder point.')
add_numbered('Select items to reorder.')
add_numbered('Click **Generate PO** to automatically create a purchase order.')

page_break()

# ===========================================================================
# SECTION 7 — INVENTORY
# ===========================================================================
add_heading('7. Inventory Management', 1)

add_heading('7.1  Inventory Snapshot', 2)
add_para('A read-only overview of current stock levels across all locations.')
add_para()
add_numbered('Click **Inventory Snapshot** from the Main Menu.')
add_numbered('Use the **status filter** to view: All, Negative, Zero, Low, or OK stock.')
add_numbered('Items expiring within 30 days are highlighted — check these regularly.')
add_para()
add_note('This screen is read-only. To adjust quantities, use Adjust Stock.')

add_heading('7.2  Adjusting Stock', 2)
add_para('Use this to correct quantities for discrepancies not related to a sale, PO, or cycle count.')
add_para()
add_numbered('Click **Adjust Stock** from the Main Menu.')
add_numbered('Search for the product and select the **location**.')
add_numbered('Enter the **adjustment quantity** (positive to add, negative to remove).')
add_numbered('Enter a **reason** for the adjustment.')
add_numbered('Click **Save**.')
add_para()
add_para('All adjustments are logged with your username, date, and reason.')

add_heading('7.3  Transferring Stock Between Locations', 2)
add_numbered('Click **Stock Transfer** from the Main Menu.')
add_numbered('Select the **source location** (where stock is coming from).')
add_numbered('Select the **destination location** (where stock is going).')
add_numbered('Search for the product and enter the **quantity to transfer**.')
add_numbered('Click **Transfer**.')

add_heading('7.4  Cycle Count (Physical Stock Count)', 2)
add_para('A cycle count is a physical inventory check where you count stock on the shelf and compare it to the system quantity.')
add_para()
add_para('**When to do a cycle count:**')
add_bullet('On a regular schedule set by management (e.g., monthly per location)')
add_bullet('When you suspect a discrepancy')
add_bullet('After a large receiving or manufacturing run')
add_para()
add_para('**How to perform a cycle count:**')
add_para()
add_numbered('Click **Cycle Count** from the Main Menu.')
add_numbered('Select the **Location** to count.')
add_numbered('Optionally tick **Show uncounted only** to focus on items not yet counted.')
add_numbered('For each product, the **System Qty** column shows the expected quantity. Enter the **Actual Qty** you physically counted.')
add_numbered('Once all counts are entered, click **Verify All** or select individual rows and click **Verify Selected**.')
add_numbered('The system records the variance and updates quantities accordingly.')
add_para()
add_note('Negative variance means less stock was found than the system shows. Investigate significant variances — they may indicate theft, damage, or recording errors.')

add_heading('7.5  Unverified Items', 2)
add_para('New products that have not yet been physically verified appear in the Unverified Items screen.')
add_numbered('Click the **Unverified Items** badge or find it in the Inventory section.')
add_numbered('Review each item.')
add_numbered('Once confirmed, click **Mark as Verified**.')

page_break()

# ===========================================================================
# SECTION 8 — PRODUCTS
# ===========================================================================
add_heading('8. Products & Parts', 1)

add_heading('8.1  Adding a New Product', 2)
add_numbered('Click **Products** from the Main Menu.')
add_numbered('Click **Add New**.')
add_numbered('Fill in the required fields:')
add_bullet('**SKU** — unique product code (required)', level=1)
add_bullet('**Name** — product description', level=1)
add_bullet('**Retail Price** — selling price to end customers', level=1)
add_bullet('**Wholesale Price** — selling price to trade customers', level=1)
add_bullet('**Reorder Point** — minimum quantity before a restock alert', level=1)
add_numbered('Fill in optional fields as needed (description, product type, attributes).')
add_numbered('Click **Save**.')

add_heading('8.2  Searching for a Product', 2)
add_numbered('Click **Products** or use Product Search from any order screen.')
add_numbered('Type part of the SKU or name.')
add_numbered('Press Enter or click **Search**.')
add_numbered('Click on a product to view its details.')
add_para()
add_note('Viewers (read-only users) can search and view products but cannot edit them.')

add_heading('8.3  Product Types & Attributes', 2)
add_bullet('**Product Types** — categories that group products (e.g., Clothing, Accessories). Managed under Products > Product Types.')
add_bullet('**Attributes** — custom options like size or colour. Managed under Products > Attribute Lists.')
add_para('Only Admins and Editors can add or change these.')

add_heading('8.4  Parts & Bills of Material (BOM)', 2)
add_para('Parts are the components used to build finished products (e.g., raw materials, packaging).')
add_bullet('**Parts Manager** — create and manage parts records.')
add_bullet('**BOM Explorer** — view and edit the Bill of Materials (what components are used to make each product, and in what quantities).')
add_para('Parts are consumed when a manufacturing work order is completed.')

page_break()

# ===========================================================================
# SECTION 9 — MANUFACTURING
# ===========================================================================
add_heading('9. Manufacturing', 1)

add_heading('9.1  Understanding Manufacturing Orders and Work Orders', 2)
add_bullet('A **Manufacturing Order (MO)** is an instruction to produce a batch of products.')
add_bullet('Each MO contains one or more **Work Orders (WOs)** — one WO per product to be made.')

add_heading('9.2  Creating a Manufacturing Order', 2)
add_numbered('Click **Manufacturing Dashboard** from the Main Menu.')
add_numbered('Click **New Manufacturing Order**.')
add_numbered('Enter a name or reference for the MO.')
add_numbered('Add each product to be manufactured: search for the product and enter the quantity to produce.')
add_numbered('Click **Save**. Work orders are automatically created for each product line.')

add_heading('9.3  Processing Work Orders', 2)
add_numbered('Click **Work Orders** from the Main Menu.')
add_numbered('Filter by date range if needed.')
add_numbered('Find the work order you are working on.')
add_numbered('Click **Mark In Progress** when production begins.')
add_numbered('When production is complete, click **Complete**.')
add_numbered('In the completion screen: confirm the **quantity produced**, enter the **cost of goods**, and select the **destination location**.')
add_numbered('Click **Confirm**.')
add_para()
add_para('The system will add the produced quantity to inventory, record the cost of goods for profit reports, and mark the work order as completed.')

page_break()

# ===========================================================================
# SECTION 10 — REPORTING
# ===========================================================================
add_heading('10. Reporting & Analytics', 1)

add_heading('10.1  KPI Dashboard', 2)
add_para('Click **KPI Dashboard** from the Main Menu to see a live summary:')
add_para()
add_table(
    ['KPI', 'What It Shows'],
    [
        ['Today\'s Orders', 'Number of orders placed today'],
        ['Today\'s Revenue', 'Total order value today'],
        ['Pending Orders', 'Orders not yet fulfilled'],
        ['In-Stock Products', 'Number of products with stock above zero'],
        ['Out of Stock / Low Stock', 'Products at zero or below reorder point'],
        ['Open Work Orders', 'Manufacturing work orders in progress'],
        ['Overdue Tasks', 'Tasks past their due date'],
        ['Total Inventory Value', 'All stock quantities × retail price'],
    ],
    col_widths=[2.5, 3.5]
)

add_heading('10.2  Reports', 2)
add_para('Click **Reports** from the Main Menu. There are five report tabs:')
add_para()

add_heading('Stock on Hand', 3)
add_para('Current inventory quantities and values across all locations. Filter by location or SKU. Values shown at retail and wholesale prices.')

add_heading('Sales by Period', 3)
add_para('All orders within a chosen date range. Filter by date, customer, or store. Totals shown in home currency.')

add_heading('COGS Summary', 3)
add_para('Cost of goods sold, pulled from completed work orders. Useful for understanding production costs over a period.')

add_heading('Cycle Count Variance', 3)
add_para('Discrepancies from past cycle counts. Helps identify recurring shrinkage or recording issues.')

add_heading('Gross Profit', 3)
add_para('Revenue minus COGS per product, with profit margin percentage. Helps identify most and least profitable products.')

add_para()
add_para('**To run any report:**')
add_numbered('Select the report tab.')
add_numbered('Set the date range and any filters.')
add_numbered('Click **Run** or **Refresh**.')
add_numbered('Click **Export to CSV** to download, or **Print** to generate a PDF.')

add_heading('10.3  Breakeven Calculator', 2)
add_numbered('Click **Breakeven Calculator** from the Main Menu.')
add_numbered('Enter your fixed costs (rent, salaries, etc.) and variable costs per unit.')
add_numbered('Enter your average selling price.')
add_numbered('The calculator shows how many units must be sold to break even.')

page_break()

# ===========================================================================
# SECTION 11 — TASKS
# ===========================================================================
add_heading('11. Task Management', 1)

add_heading('11.1  Creating a Task', 2)
add_numbered('Click **Task Manager** from the Main Menu.')
add_numbered('Click **New Task**.')
add_numbered('Enter the **Title**, **Assigned To** (select a user), **Due Date**, and **Details**.')
add_numbered('Click **Save**.')

add_heading('11.2  Viewing and Managing Tasks', 2)
add_bullet('Use the **filter by user** dropdown to see tasks assigned to a specific person.')
add_bullet('Overdue tasks are highlighted in red.')
add_bullet('To mark a task done: select it and click **Mark Done**, or select multiple tasks and use **Bulk Mark Done**.')
add_bullet('Click on a task to open the Task Detail screen and view the full history and comments.')

add_heading('11.3  Task Comments & Mentions', 2)
add_para('Inside a task\'s detail screen, you can add comments to discuss the task.')
add_bullet('Type your comment in the comment box.')
add_bullet('To notify a specific user, type **@username** in your comment (e.g., @jane). They will receive a badge notification on their Main Menu.')
add_bullet('Click **Add Comment** to post.')
add_para()
add_note('Check your mention badges on the Main Menu regularly to stay up to date on tasks where you have been mentioned.')

add_heading('11.4  Task Statuses', 2)
add_table(
    ['Status', 'Meaning'],
    [
        ['Open', 'Task created, not yet started'],
        ['In Progress', 'Work has begun'],
        ['Done', 'Task is complete'],
    ],
    col_widths=[2.0, 4.0]
)

page_break()

# ===========================================================================
# SECTION 12 — ADMINISTRATION
# ===========================================================================
add_heading('12. System Administration', 1)
add_note('The following sections apply to Administrators only.')
add_para()

add_heading('12.1  Managing Users', 2)
add_numbered('Click **User Management** from the Main Menu.')
add_numbered('To **add a new user**: click New User, enter their name, username, password, and assign their role and permissions. Click Save.')
add_numbered('To **edit a user**: select them from the list and update their details or permissions. Click Save.')
add_numbered('To **deactivate a user**: edit their record and disable their account.')
add_para()
add_note('Create accounts individually — do not share login credentials. Every action is logged against the user account.')

add_heading('12.2  System Settings', 2)
add_para('Click **Settings** from the Main Menu to configure:')
add_para()
add_table(
    ['Setting', 'Description'],
    [
        ['Company Logo', 'Upload the logo shown in the application'],
        ['Currency', 'Set home currency and exchange rates'],
        ['Password Policy', 'Set max failed login attempts and lockout duration'],
        ['Backup Schedule', 'How often the database is backed up'],
        ['Backup Folder', 'Destination folder for backup files'],
        ['Accent Colour', 'Customise the application colour scheme'],
        ['SMTP Email', 'Configure email settings for notifications'],
        ['Admin Contact', 'Admin name/phone shown on the login screen'],
        ['Remember Username', 'Enable/disable the remember username option'],
    ],
    col_widths=[2.0, 4.0]
)

add_heading('12.3  Viewing Logs', 2)
add_bullet('**Activity Log** — a record of all data changes (creates, edits, deletes) by all users. Use this to investigate discrepancies.')
add_bullet('**Login Log** — all login and logout events with timestamps. Use this to verify who accessed the system and when.')
add_para()
add_para('To access logs: click Activity Log or Login Log from the Main Menu (requires Log permission or Admin role).')

page_break()

# ===========================================================================
# SECTION 13 — IMPORT / EXPORT
# ===========================================================================
add_heading('13. Data Import & Export', 1)

add_heading('13.1  Importing Data', 2)
add_numbered('Click **Imports** from the Main Menu.')
add_numbered('Select the data type to import (e.g., Products, Customers, Inventory).')
add_numbered('Download the **CSV template** to see the required format.')
add_numbered('Prepare your CSV file following the template exactly.')
add_numbered('Click **Browse**, select your file, then click **Import**.')
add_numbered('Review any errors or warnings shown after the import.')
add_para()
add_note('Always review the template before creating your import file. Test with a small batch first.')

add_heading('13.2  Exporting Data', 2)
add_numbered('Click **Exports** from the Main Menu.')
add_numbered('Select the data type to export.')
add_numbered('Apply any filters (e.g., date range).')
add_numbered('Click **Export to CSV** and choose where to save the file.')
add_para()
add_para('Reports also have their own export buttons — see Section 10.')

page_break()

# ===========================================================================
# SECTION 14 — TROUBLESHOOTING
# ===========================================================================
add_heading('14. Troubleshooting & FAQs', 1)

faqs = [
    (
        'My account is locked.',
        'Contact your administrator. Accounts lock after repeated failed login attempts. The lockout clears automatically after the configured period, or an admin can reset it manually.'
    ),
    (
        'I can\'t see a button or screen.',
        'Your role or permissions may not include access to that area. Contact your administrator if you believe you should have access.'
    ),
    (
        'The system logged me out automatically.',
        'The session timeout is 30 minutes of inactivity. Log back in. If you were in the middle of data entry that was not saved, you will need to re-enter it.'
    ),
    (
        'Stock quantities look wrong.',
        '1. Check the Inventory Snapshot to confirm current quantities.\n2. Check the Activity Log to see if a recent adjustment was made.\n3. Run a Cycle Count on the affected location to verify physical quantities.\n4. If the discrepancy cannot be explained, use Adjust Stock with a written reason.'
    ),
    (
        'A purchase order is stuck on "Partially Received".',
        'This means not all items have been received yet. When the remaining items arrive, open the PO, click Receive Items, and complete the receipt. The PO status updates to Received once all quantities are confirmed.'
    ),
    (
        'A Shopify order is not appearing in the Sales Dashboard.',
        '1. Click Fetch Latest to pull fresh data from Shopify.\n2. Check your date range filter — the order may fall outside the selected period.\n3. Check the minimum amount filter is not hiding the order.\n4. If the order still does not appear, contact your administrator to check the Shopify sync configuration.'
    ),
    (
        'How do I correct a mistake on an order?',
        'Open the order and edit it. All changes are recorded in the Activity Log. If an order cannot be edited (e.g., it is finalised), contact your administrator.'
    ),
    (
        'I cannot find a product when searching.',
        'Try searching by SKU instead of name, or vice versa. Confirm the product exists in the Products screen. The product may have a different name or SKU than expected.'
    ),
]

for question, answer in faqs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(question)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)
    add_para(answer)
    add_para()

# ---------------------------------------------------------------------------
# Footer note
# ---------------------------------------------------------------------------
doc.add_paragraph()
hr = doc.add_paragraph()
hr.paragraph_format.space_before = Pt(12)
run = hr.add_run('For technical issues or system errors, contact your system administrator.')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

footer_note = doc.add_paragraph()
run2 = footer_note.add_run('This document should be reviewed and updated whenever significant changes are made to the software.')
run2.italic = True
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(128, 128, 128)

# ===========================================================================
# SAVE
# ===========================================================================
output_path = r'E:\Onedrive - Jvnction\Documents - VanGo Production\Production Tools\Production 2023\Production Tools 2023\New System 2026\JaneERP\docs\JaneERP_SOP.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
